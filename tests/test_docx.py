"""Markdown → DOCX 변환 및 /export/docx 엔드포인트 테스트."""
from __future__ import annotations

import io

from docx import Document
from fastapi.testclient import TestClient

from app.main import app
from app.services import docx_export

MD = """# 테스트 기획서

## 프로젝트 개요
**핵심** 요약 문장.

## PESTEL 분석
| 요인 | 주요 내용 | 기회 | 위협 | 대응 |
|---|---|---|---|---|
| Political | 규제 | 지원 | 불확실 | 모니터링 |
| Economic | 성장 | 수요 | 비용 | 효율화 |

## 참고자료
- 출처 A — https://a.io
- https://b.io
"""


def test_build_docx_structure():
    doc = docx_export.build_docx(MD)
    heads = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert "테스트 기획서" in heads
    assert "PESTEL 분석" in heads
    # 표 1개, 헤더 5열 + 데이터 2행
    assert len(doc.tables) == 1
    t = doc.tables[0]
    assert len(t.columns) == 5 and len(t.rows) == 3
    assert t.rows[0].cells[0].text == "요인"
    assert t.rows[1].cells[0].text == "Political"


def test_docx_bytes_is_valid_zip_openable():
    data = docx_export.docx_bytes(MD)
    assert data[:2] == b"PK"                      # docx = zip 컨테이너
    Document(io.BytesIO(data))                     # 다시 열려야 유효


def test_export_docx_endpoint():
    c = TestClient(app)
    r = c.post("/export/docx", json={"project_name": "테스트", "markdown": MD})
    assert r.status_code == 200
    assert "wordprocessingml" in r.headers["content-type"]
    assert r.content[:2] == b"PK"
