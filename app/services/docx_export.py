"""최종 기획서(Markdown) → Word(.docx) 변환.

최종본은 `# 제목`, `## 섹션`, 마크다운 표(PESTEL 등), 목록(-), 굵게(**)를 쓴다.
python-docx로 이 요소들을 실제 Word 서식으로 렌더한다(외부 변환기 불필요).
"""
from __future__ import annotations

import io
import re
from pathlib import Path

from docx import Document

OUTPUT_DIR = Path("outputs")
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_TABLE_ROW = re.compile(r"^\s*\|(.+)\|\s*$")
_TABLE_SEP = re.compile(r"^\s*\|[\s:|-]+\|\s*$")


def _slugify(name: str) -> str:
    slug = re.sub(r"[^0-9A-Za-z가-힣]+", "-", name).strip("-")
    return slug or "plan"


def _cells(line: str) -> list[str]:
    return [c.strip() for c in _TABLE_ROW.match(line).group(1).split("|")]


def _add_runs(paragraph, text: str) -> None:
    """**굵게** 를 인식해 run 을 나눠 추가한다."""
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        paragraph.add_run(m.group(1)).bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def build_docx(markdown: str) -> Document:
    doc = Document()
    lines = (markdown or "").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        # 표
        if _TABLE_ROW.match(line) and i + 1 < len(lines) and _TABLE_SEP.match(lines[i + 1]):
            header = _cells(line)
            i += 2
            rows = []
            while i < len(lines) and _TABLE_ROW.match(lines[i]) and not _TABLE_SEP.match(lines[i]):
                rows.append(_cells(lines[i]))
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Table Grid"
            for j, h in enumerate(header):
                cell = table.rows[0].cells[j]
                cell.paragraphs[0].add_run(h).bold = True
            for r in rows:
                cells = table.add_row().cells
                for j in range(len(header)):
                    cells[j].text = r[j] if j < len(r) else ""
            doc.add_paragraph()
            continue
        # 제목
        m = re.match(r"^\s*(#{1,3})\s+(.*)", line)
        if m:
            doc.add_heading(m.group(2).strip(), level=len(m.group(1)))
            i += 1
            continue
        # 목록
        if re.match(r"^\s*-\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, re.sub(r"^\s*-\s+", "", line))
            i += 1
            continue
        # 문단 / 빈 줄
        if line.strip():
            _add_runs(doc.add_paragraph(), line.strip())
        i += 1
    return doc


def docx_bytes(markdown: str) -> bytes:
    buf = io.BytesIO()
    build_docx(markdown).save(buf)
    return buf.getvalue()


def save_docx(project_name: str, markdown: str) -> str:
    OUTPUT_DIR.mkdir(exist_ok=True)
    path = OUTPUT_DIR / f"{_slugify(project_name)}.docx"
    build_docx(markdown).save(str(path))
    return str(path)
